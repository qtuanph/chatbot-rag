from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re

import fitz
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.services.ocr import get_ocr_service
from app.core.config import settings


@dataclass
class IngestedNode:
    ref: str
    heading: str
    full_text: str
    summary: str | None = None
    page_range: str | None = None
    level: int = 0
    order_index: int = 0
    parent_ref: str | None = "root"


@dataclass
class IngestionArtifact:
    node_count: int
    non_empty_node_count: int
    total_text_chars: int
    avg_text_chars: float
    max_level: int
    extraction_mode: str
    content_coverage_ratio: float
    warnings: list[str]
    node_type_counts: dict[str, int]

    def to_dict(self) -> dict:
        return {
            "node_count": self.node_count,
            "non_empty_node_count": self.non_empty_node_count,
            "total_text_chars": self.total_text_chars,
            "avg_text_chars": round(self.avg_text_chars, 2),
            "max_level": self.max_level,
            "extraction_mode": self.extraction_mode,
            "content_coverage_ratio": round(self.content_coverage_ratio, 4),
            "warnings": self.warnings,
            "node_type_counts": self.node_type_counts,
        }


def extract_nodes(filename: str, content: bytes) -> list[IngestedNode]:
    nodes, _ = extract_nodes_with_artifact(filename, content)
    return nodes


def extract_nodes_with_artifact(filename: str, content: bytes) -> tuple[list[IngestedNode], IngestionArtifact]:
    suffix = Path(filename).suffix.lower()
    extraction_mode = "text"

    if suffix == ".pdf":
        nodes = _extract_pdf_nodes(content)
        extraction_mode = "pdf"
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
        nodes = [_ocr_image_node(filename, content)]
        extraction_mode = "ocr-image"
    elif suffix == ".docx":
        nodes = _extract_docx_nodes(content)
        extraction_mode = "docx"
    elif suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        nodes = _extract_xlsx_nodes(content)
        extraction_mode = "xlsx"
    elif suffix == ".md":
        nodes = _extract_markdown_nodes(content)
        extraction_mode = "markdown"
    else:
        nodes = _extract_plain_text_nodes(filename, content)
        extraction_mode = "plain"

    artifact = _build_ingestion_artifact(nodes, extraction_mode)
    return nodes, artifact


def _build_ingestion_artifact(nodes: list[IngestedNode], extraction_mode: str) -> IngestionArtifact:
    node_count = len(nodes)
    non_empty_nodes = [node for node in nodes if node.full_text and node.full_text.strip()]
    non_empty_node_count = len(non_empty_nodes)
    total_text_chars = sum(len(node.full_text.strip()) for node in non_empty_nodes)
    avg_text_chars = (total_text_chars / non_empty_node_count) if non_empty_node_count else 0.0
    max_level = max((node.level for node in nodes), default=0)

    coverage_ratio = (non_empty_node_count / node_count) if node_count else 0.0
    warnings: list[str] = []
    if node_count == 0:
        warnings.append("no_nodes_extracted")
    if non_empty_node_count == 0 and node_count > 0:
        warnings.append("all_nodes_empty")
    if coverage_ratio < 0.5 and node_count > 0:
        warnings.append("low_content_coverage")
    if total_text_chars < 120 and node_count > 0:
        warnings.append("very_low_total_text")

    node_type_counts: dict[str, int] = {}
    for node in nodes:
        prefix = node.ref.split(":", 1)[0] if ":" in node.ref else "generic"
        node_type_counts[prefix] = node_type_counts.get(prefix, 0) + 1

    return IngestionArtifact(
        node_count=node_count,
        non_empty_node_count=non_empty_node_count,
        total_text_chars=total_text_chars,
        avg_text_chars=avg_text_chars,
        max_level=max_level,
        extraction_mode=extraction_mode,
        content_coverage_ratio=coverage_ratio,
        warnings=warnings,
        node_type_counts=node_type_counts,
    )


def _extract_pdf_nodes(content: bytes) -> list[IngestedNode]:
    if settings.ocr_strategy in {"markdown", "hybrid"}:
        markdown = get_ocr_service().document_to_markdown(content, filetype="pdf")
        if markdown:
            md_nodes = _extract_markdown_nodes(markdown.encode("utf-8"))
            if _has_usable_nodes(md_nodes):
                return md_nodes

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        nodes: list[IngestedNode] = []
        order_index = 1
        for index, page in enumerate(doc, start=1):
            text = _normalize_text(page.get_text("text"))
            if not text:
                text = _normalize_text(_ocr_pdf_page(page))
            if not text:
                continue

            sections = _split_into_sections(text)
            if len(sections) == 1:
                body = sections[0]
                nodes.append(
                    IngestedNode(
                        ref=f"page:{index}",
                        heading=f"Page {index}",
                        full_text=body,
                        summary=body[:240],
                        page_range=str(index),
                        level=1,
                        order_index=order_index,
                        parent_ref="root",
                    )
                )
                order_index += 1
                continue

            page_heading = f"Page {index}"
            nodes.append(
                IngestedNode(
                    ref=f"page:{index}",
                    heading=page_heading,
                    full_text=text,
                    summary=text[:240],
                    page_range=str(index),
                    level=1,
                    order_index=order_index,
                    parent_ref="root",
                )
            )
            order_index += 1
            for section_index, section in enumerate(sections, start=1):
                nodes.append(
                    IngestedNode(
                        ref=f"section:{index}:{section_index}",
                        heading=f"Section {index}.{section_index}",
                        full_text=section,
                        summary=section[:240],
                        page_range=str(index),
                        level=2,
                        order_index=order_index,
                        parent_ref=f"page:{index}",
                    )
                )
                order_index += 1

        toc_nodes = _extract_pdf_toc_nodes(doc, nodes, order_index)
        nodes.extend(toc_nodes)

        return nodes
    finally:
        doc.close()


def _has_usable_nodes(nodes: list[IngestedNode]) -> bool:
    if not nodes:
        return False
    non_empty = [n for n in nodes if n.full_text and n.full_text.strip()]
    total_chars = sum(len(n.full_text.strip()) for n in non_empty)
    return len(non_empty) >= 2 and total_chars >= 200


def _extract_pdf_toc_nodes(doc, existing_nodes: list[IngestedNode], start_order: int) -> list[IngestedNode]:
    try:
        toc = doc.get_toc(simple=True)
    except Exception:
        return []

    if not toc:
        return []

    page_to_ref: dict[int, str] = {}
    for node in existing_nodes:
        if node.ref.startswith("page:"):
            try:
                page_no = int(node.ref.split(":", 1)[1])
                page_to_ref[page_no] = node.ref
            except Exception:
                continue

    level_last_ref: dict[int, str] = {}
    toc_nodes: list[IngestedNode] = []
    order_index = start_order
    for idx, item in enumerate(toc, start=1):
        if len(item) < 3:
            continue
        level, title, page = item[0], str(item[1]).strip(), int(item[2])
        if not title:
            continue
        parent_ref = "root"
        if level > 1 and (level - 1) in level_last_ref:
            parent_ref = level_last_ref[level - 1]
        elif page in page_to_ref:
            parent_ref = page_to_ref[page]

        ref = f"toc:{idx}"
        level_last_ref[level] = ref
        toc_nodes.append(
            IngestedNode(
                ref=ref,
                heading=title,
                full_text=title,
                summary=title,
                page_range=str(page),
                level=max(1, min(level + 1, 6)),
                order_index=order_index,
                parent_ref=parent_ref,
            )
        )
        order_index += 1

    return toc_nodes


def _ocr_pdf_page(page) -> str:
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    return get_ocr_service().image_to_text(pix.tobytes("png"))


def _ocr_image_node(filename: str, content: bytes) -> IngestedNode:
    text = _normalize_text(get_ocr_service().image_to_text(content))
    return IngestedNode(
        ref=f"image:{Path(filename).name}",
        heading=Path(filename).name,
        full_text=text,
        summary=text[:240] if text else None,
        level=1,
        parent_ref="root",
    )


def _extract_docx_nodes(content: bytes) -> list[IngestedNode]:
    document = DocxDocument(BytesIO(content))
    nodes: list[IngestedNode] = []
    order_index = 1
    current_heading = "Document"
    current_ref = "docx:root"
    current_level = 0
    heading_count = 1
    buffer: list[str] = []
    last_was_blank = True
    level_stack: dict[int, str] = {0: "root"}

    def flush_chunk() -> None:
        nonlocal order_index
        if not buffer:
            return
        body = _normalize_text("\n".join(buffer))
        if not body:
            return
        heading = current_heading if current_heading else "Section"
        nodes.append(
            IngestedNode(
                ref=f"docx:{order_index}",
                heading=heading,
                full_text=body,
                summary=body[:240],
                level=max(1, current_level + 1) if heading != "Document" else 1,
                order_index=order_index,
                parent_ref=current_ref,
            )
        )
        order_index += 1

    def next_heading_ref() -> str:
        nonlocal heading_count
        ref = f"docx-heading:{heading_count}"
        heading_count += 1
        return ref

    for p in document.paragraphs:
        text = _normalize_text(p.text)
        if not text:
            last_was_blank = True
            continue
        style_name = (p.style.name or "").lower()
        is_heading = style_name.startswith("heading") or _looks_like_manual_heading(text, last_was_blank, p)
        if is_heading:
            flush_chunk()
            heading_level = _heading_level_from_style(style_name, text)
            current_heading = text
            current_level = heading_level
            current_ref = next_heading_ref()
            level_stack[heading_level] = current_ref
            for deeper in range(heading_level + 1, 8):
                level_stack.pop(deeper, None)
            parent_ref = level_stack.get(max(heading_level - 1, 0), "root")
            nodes.append(
                IngestedNode(
                    ref=current_ref,
                    heading=text,
                    full_text=text,
                    summary=text[:240],
                    level=heading_level,
                    order_index=order_index,
                    parent_ref=parent_ref,
                )
            )
            order_index += 1
            buffer = []
            last_was_blank = False
            continue
        buffer.append(text)
        last_was_blank = False

    flush_chunk()
    if not nodes:
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        return _make_chunk_nodes("Document", "\n".join(paragraphs), target_chars=1600)
    return nodes


def _extract_xlsx_nodes(content: bytes) -> list[IngestedNode]:
    workbook = load_workbook(BytesIO(content), data_only=True)
    nodes: list[IngestedNode] = []
    order_index = 1
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell).strip() for cell in row]
            line = " | ".join(value for value in values if value)
            if line:
                rows.append(line)
        if rows:
            text = "\n".join(rows)
            nodes.append(
                IngestedNode(
                    ref=f"sheet:{order_index}",
                    heading=sheet.title,
                    full_text=text,
                    summary=text[:240],
                    level=1,
                    order_index=order_index,
                    parent_ref="root",
                )
            )
            order_index += 1
    return nodes


def _extract_markdown_nodes(content: bytes) -> list[IngestedNode]:
    text = _normalize_text(content.decode("utf-8", errors="ignore"))
    if not text:
        return []

    lines = text.splitlines()
    nodes: list[IngestedNode] = []
    order_index = 1
    section_buffer: list[str] = []
    section_heading = "Document"
    section_level = 1
    section_ref = "md:root"
    level_stack: dict[int, str] = {0: "root"}

    def flush_section() -> None:
        nonlocal order_index
        if not section_buffer:
            return
        body = _normalize_text("\n".join(section_buffer))
        if not body:
            return
        parent_level = max(section_level - 1, 0)
        parent_ref = level_stack.get(parent_level, "root")
        nodes.append(
            IngestedNode(
                ref=section_ref,
                heading=section_heading,
                full_text=body,
                summary=body[:240],
                level=max(1, section_level),
                order_index=order_index,
                parent_ref=parent_ref,
            )
        )
        order_index += 1

    section_counter = 1
    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if heading_match:
            flush_section()
            hashes, title = heading_match.groups()
            level = len(hashes)
            section_heading = title.strip()
            section_level = level
            section_ref = f"md:{section_counter}"
            section_counter += 1
            level_stack[level] = section_ref
            for deeper in range(level + 1, 8):
                level_stack.pop(deeper, None)
            section_buffer = [section_heading]
            continue
        section_buffer.append(line)

    flush_section()
    if not nodes:
        return _make_chunk_nodes("Document", text, target_chars=1600)
    return nodes


def _extract_plain_text_nodes(filename: str, content: bytes) -> list[IngestedNode]:
    text = _normalize_text(content.decode("utf-8", errors="ignore"))
    if not text:
        return []
    md_like = _extract_markdown_nodes(content)
    if md_like:
        return md_like
    if any(_looks_like_heading(line) for line in text.splitlines()):
        return _extract_heading_aware_plain_text_nodes(Path(filename).name, text)
    return _make_chunk_nodes(Path(filename).name, text, target_chars=1600)


def _extract_heading_aware_plain_text_nodes(filename: str, text: str) -> list[IngestedNode]:
    lines = [line.strip() for line in text.splitlines()]
    nodes: list[IngestedNode] = []
    order_index = 1
    buffer: list[str] = []
    current_heading = filename
    current_ref = "txt:root"
    section_counter = 1

    def flush() -> None:
        nonlocal order_index
        if not buffer:
            return
        body = _normalize_text("\n".join(buffer))
        if not body:
            return
        nodes.append(
            IngestedNode(
                ref=f"txt:{order_index}",
                heading=current_heading,
                full_text=body,
                summary=body[:240],
                level=1,
                order_index=order_index,
                parent_ref=current_ref,
            )
        )
        order_index += 1

    for idx, line in enumerate(lines):
        if _looks_like_heading(line):
            flush()
            current_heading = line
            current_ref = f"txt-heading:{section_counter}"
            section_counter += 1
            nodes.append(
                IngestedNode(
                    ref=current_ref,
                    heading=line,
                    full_text=line,
                    summary=line[:240],
                    level=1,
                    order_index=order_index,
                    parent_ref="root",
                )
            )
            order_index += 1
            buffer = []
            continue
        if line:
            buffer.append(line)

    flush()
    return nodes if nodes else _make_chunk_nodes(filename, text, target_chars=1600)


def _split_into_sections(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return []

    sections: list[str] = []
    buffer: list[str] = []
    for line in lines:
        if _looks_like_heading(line) and buffer and len(" ".join(buffer)) >= 80:
            sections.append("\n".join(buffer).strip())
            buffer = [line]
        else:
            buffer.append(line)
    if buffer:
        sections.append("\n".join(buffer).strip())

    return [section for section in sections if section]


def _looks_like_heading(line: str) -> bool:
    if len(line) > 120 or len(line) < 3:
        return False
    if line.endswith(":"):
        return True
    if re.match(r"^(CHAPTER|CHƯƠNG|MỤC|PHẦN|SECTION|ARTICLE)\b", line, flags=re.IGNORECASE):
        return True
    if re.match(r"^(\d+(?:\.\d+){0,3})\s+\S+", line):
        return True
    alpha = [ch for ch in line if ch.isalpha()]
    if alpha:
        upper_ratio = sum(ch.isupper() for ch in alpha) / len(alpha)
        if upper_ratio >= 0.85 and len(line.split()) <= 12:
            return True
    if len(line.split()) <= 10 and line[0].isupper() and not line.endswith("."):
        return True
    return False


def _looks_like_manual_heading(text: str, last_was_blank: bool, paragraph) -> bool:
    if _looks_like_heading(text):
        return True
    if not last_was_blank and len(text.split()) > 10:
        return False
    if len(text) > 100:
        return False
    if not text[:1].isalpha():
        return False
    if text.endswith("."):
        return False

    bold_runs = 0
    total_runs = 0
    try:
        for run in paragraph.runs:
            total_runs += 1
            if run.bold:
                bold_runs += 1
    except Exception:
        total_runs = 0

    if total_runs and bold_runs / total_runs >= 0.5:
        return True
    if re.match(r"^(\d+|[IVXLC]+)[\.)]\s+\S+", text):
        return True
    if len(text.split()) <= 8 and text[0].isupper():
        return True
    return False


def _heading_level_from_style(style_name: str, text: str) -> int:
    match = re.search(r"heading\s*(\d+)", style_name)
    if match:
        return max(1, min(int(match.group(1)), 6))
    if re.match(r"^(CHAPTER|CHƯƠNG|MỤC|PHẦN)\b", text, flags=re.IGNORECASE):
        return 1
    if re.match(r"^\d+\.\d+", text):
        return 2
    return 1


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _make_chunk_nodes(heading: str, text: str, target_chars: int) -> list[IngestedNode]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return []

    nodes: list[IngestedNode] = []
    chunk: list[str] = []
    chunk_size = 0
    order_index = 1
    label_prefix = heading if heading != "Document" else "Section"

    for line in lines:
        chunk.append(line)
        chunk_size += len(line) + 1
        if chunk_size >= target_chars:
            body = "\n".join(chunk)
            nodes.append(
                IngestedNode(
                    ref=f"chunk:{order_index}",
                    heading=f"{label_prefix} {order_index}",
                    full_text=body,
                    summary=body[:240],
                    level=1,
                    order_index=order_index,
                    parent_ref="root",
                )
            )
            order_index += 1
            chunk = []
            chunk_size = 0

    if chunk:
        body = "\n".join(chunk)
        nodes.append(
            IngestedNode(
                ref=f"chunk:{order_index}",
                heading=f"{label_prefix} {order_index}",
                full_text=body,
                summary=body[:240],
                level=1,
                order_index=order_index,
                parent_ref="root",
            )
        )

    return nodes
