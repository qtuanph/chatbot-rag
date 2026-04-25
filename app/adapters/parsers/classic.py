"""
Classic Parser: Fallback parser for text, Markdown, DOCX, XLSX, and image-only content.
Used when Docling parser is unavailable or fails.
"""

import logging
import uuid
from typing import List, Tuple
from io import BytesIO

from app.adapters.base import (
    BaseParser,
    IngestedNode,
    ParsingMetadata,
    ParsedNodeType,
)
from app.core.file_formats import extract_file_format
from app.core.exceptions import ParsingException

logger = logging.getLogger(__name__)


class ClassicParser(BaseParser):
    """
    Fallback parser: Handles plain text, Markdown, DOCX, XLSX formats.
    Simple, non-hierarchical extraction; suitable for emergency fallback.
    """

    def parse(
        self,
        filename: str,
        content: bytes,
    ) -> Tuple[List[IngestedNode], ParsingMetadata]:
        """
        Parse document using format-specific fallback extractors.
        
        Args:
            filename: Document filename (for format detection)
            content: Raw file bytes
        
        Returns:
            Tuple of (IngestedNode list, ParsingMetadata)
        
        Raises:
            ParsingException: If parsing fails
        """
        import time
        start_time = time.time()
        
        source_format = extract_file_format(filename)
        
        try:
            nodes = []
            
            # Dispatch to format-specific extractor
            if source_format in {'pdf', 'image'}:
                # For PDFs without Docling, attempt simple page-based extraction
                nodes = self._extract_pdf_pages(filename, content)
            elif source_format == 'docx':
                nodes = self._extract_docx_text(filename, content)
            elif source_format == 'xlsx':
                nodes = self._extract_xlsx_sheets(filename, content)
            elif source_format in {'markdown', 'text'}:
                nodes = self._extract_text_markdown(filename, content)
            else:
                # Try plain text as last resort
                nodes = self._extract_text_markdown(filename, content)
            
            if not nodes:
                raise ParsingException(
                    f"No content extracted from {filename} (format: {source_format})",
                    error_code="CLASSIC_PARSER_NO_CONTENT",
                    details={'filename': filename, 'source_format': source_format}
                )
            
            parse_time_ms = (time.time() - start_time) * 1000
            quality_score = 0.7  # Lower than Docling direct extraction
            
            metadata = ParsingMetadata(
                engine_used="classic",
                source_format=source_format,
                docling_used=False,
                fallback_used=True,
                quality_score=quality_score,
                parse_time_ms=parse_time_ms,
                node_count=len(nodes),
                total_text_chars=sum(len(node.text) for node in nodes),
                sections_data=self._build_sections_data(nodes),
                warnings=[f"Using fallback parser (not Docling); quality may be lower"],
            )
            
            logger.info(f"Parsed {filename} with classic parser: {len(nodes)} nodes (format: {source_format})")
            return nodes, metadata
        
        except Exception as e:
            raise ParsingException(
                f"Classic parser failed for {filename}: {str(e)}",
                error_code="CLASSIC_PARSER_FAILED",
                details={'filename': filename, 'source_format': source_format, 'error': str(e)}
            )

    def _extract_pdf_pages(self, filename: str, content: bytes) -> List[IngestedNode]:
        """Extract text from PDF on a page-by-page basis (simple fallback)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed; cannot extract PDF pages")
            return []

        nodes = []
        try:
            reader = PdfReader(BytesIO(content))
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text.strip():
                    node = IngestedNode(
                        node_id=str(uuid.uuid4()),
                        document_id=filename,
                        text=text,
                        node_type=ParsedNodeType.PAGE,
                        page_number=page_num,
                        order=page_num - 1,
                        metadata={'source_format': 'pdf', 'page_count': len(reader.pages)},
                    )
                    nodes.append(node)
        except Exception as e:
            logger.warning(f"PDF page extraction failed: {str(e)}")
        
        return nodes

    def _extract_docx_text(self, filename: str, content: bytes) -> List[IngestedNode]:
        """Extract text from DOCX files."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed; cannot extract DOCX text")
            return []
        
        nodes = []
        try:
            doc = DocxDocument(BytesIO(content))
            text_buffer = []
            
            for para_idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_buffer.append(para.text)
                    
                    # Create node per paragraph or every N paragraphs
                    if len(text_buffer) >= 5 or para_idx == len(doc.paragraphs) - 1:
                        node = IngestedNode(
                            node_id=str(uuid.uuid4()),
                            document_id=filename,
                            text='\n'.join(text_buffer),
                            node_type=ParsedNodeType.PARAGRAPH,
                            order=para_idx,
                            metadata={'source_format': 'docx', 'paragraph_count': len(doc.paragraphs)},
                        )
                        nodes.append(node)
                        text_buffer = []
            
            # Handle tables
            for table_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    rows.append(row_text)
                
                table_text = '\n'.join(rows)
                if table_text.strip():
                    node = IngestedNode(
                        node_id=str(uuid.uuid4()),
                        document_id=filename,
                        text=table_text,
                        node_type=ParsedNodeType.TABLE,
                        order=len(nodes),
                        metadata={'source_format': 'docx', 'table_index': table_idx},
                    )
                    nodes.append(node)
        
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {str(e)}")
        
        return nodes

    def _extract_xlsx_sheets(self, filename: str, content: bytes) -> List[IngestedNode]:
        """Extract text from XLSX files (sheet by sheet)."""
        try:
            import openpyxl
        except ImportError:
            logger.warning("openpyxl not installed; cannot extract XLSX text")
            return []
        
        nodes = []
        try:
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
            
            for sheet_idx, sheet in enumerate(workbook.sheetnames):
                ws = workbook[sheet]
                rows = []
                
                for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        rows.append(row_text)
                
                sheet_text = '\n'.join(rows)
                if sheet_text.strip():
                    node = IngestedNode(
                        node_id=str(uuid.uuid4()),
                        document_id=filename,
                        text=sheet_text,
                        node_type=ParsedNodeType.TABLE,
                        order=sheet_idx,
                        metadata={'source_format': 'xlsx', 'sheet_name': sheet, 'sheet_index': sheet_idx},
                    )
                    nodes.append(node)
        
        except Exception as e:
            logger.warning(f"XLSX extraction failed: {str(e)}")
        
        return nodes

    def _extract_text_markdown(self, filename: str, content: bytes) -> List[IngestedNode]:
        """Extract text from plain text or Markdown files."""
        try:
            text = content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.warning(f"Failed to decode text file {filename}: {str(e)}")
            return []
        
        nodes = []
        if text.strip():
            # Simple paragraph-based splitting
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            for para_idx, para in enumerate(paragraphs):
                node = IngestedNode(
                    node_id=str(uuid.uuid4()),
                    document_id=filename,
                    text=para,
                    node_type=ParsedNodeType.PARAGRAPH,
                    order=para_idx,
                    metadata={'source_format': 'text' if filename.endswith('.txt') else 'markdown'},
                )
                nodes.append(node)

        return nodes

    def _build_sections_data(self, nodes: List[IngestedNode]) -> List[dict]:
        """Generate flat section records from classic parser nodes.

        Ensures the tree/detail page shows content even when Docling fails.
        """
        sections = []
        for i, node in enumerate(nodes):
            page_label = f"Trang {node.page_number}" if node.page_number else f"Phần {i + 1}"
            sections.append({
                "section_id": f"classic-{node.node_id}",
                "parent_section_id": None,
                "title": node.section_title or page_label,
                "content": node.text,
                "section_type": node.node_type.value if hasattr(node.node_type, 'value') else 'section',
                "level": 1,
                "order_index": i,
                "page_range": str(node.page_number) if node.page_number else None,
            })
        return sections
